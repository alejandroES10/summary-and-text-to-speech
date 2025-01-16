
from fastapi import FastAPI, File, HTTPException, UploadFile, BackgroundTasks
from fastapi.responses import FileResponse, JSONResponse
from .audio_processor import text_to_mp3
from fastapi.middleware.cors import CORSMiddleware
import os
from pydantic import BaseModel
from docx import Document
from .document_processor import process_document
from fastapi.responses import FileResponse
import shutil

class TextoRequest(BaseModel):
    text: str


app = FastAPI()


origins = [
    "*", 
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,  
    allow_credentials=True,
    allow_methods=["*"], 
    allow_headers=["*"],  
)


@app.post("/convert-text-to-audio/")
async def convert_text_to_audio(request: TextoRequest):
    """
    Convierte el texto recibido en un archivo MP3 temporal y lo devuelve.
    """
    file_mp3 = text_to_mp3(request.text)
    
    if file_mp3:
        # Devolver el archivo como respuesta
        response = FileResponse(file_mp3, media_type='audio/mpeg', headers={"Content-Disposition": "attachment; filename=audio_resumen.mp3"})
        
        return response
    else:
        raise HTTPException(status_code=500, detail="No se pudo generar el archivo de audio.")
    
    
#********************************** ********************* ******************

def remove_file(file_path: str):
    """Elimina un archivo si existe."""
    if os.path.exists(file_path):
        os.remove(file_path)

@app.post("/summarize/")
async def process_document_endpoint(
    file: UploadFile = File(...), background_tasks: BackgroundTasks = BackgroundTasks()
):
    temp_file_path = f"temp_{file.filename}"
    word_file_path = f"Resumen_{file.filename}"

    try:
        # Guardar el archivo subido como temporal
        with open(temp_file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        # Procesar el archivo y generar un resumen
        result = await process_document(temp_file_path)
        final_summary = result['generate_final_summary']['final_summary']
        summary = final_summary.replace('*','')
        
        # Crear un documento Word con el resumen
        doc = Document()
        doc.add_heading('Resumen del Documento', 0)
        doc.add_paragraph(summary)

        # Guardar el archivo Word
        doc.save(word_file_path)

        # Programar la eliminación de los archivos temporales después de enviar la respuesta
        background_tasks.add_task(remove_file, temp_file_path)
        background_tasks.add_task(remove_file, word_file_path)

        # Devolver el archivo Word como respuesta
        return FileResponse(
            word_file_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=word_file_path,
            background=background_tasks  # Tareas en segundo plano
        )

    except Exception as e:
        # Capturar y manejar el error
        error_message = f"Hubo un error al procesar el archivo: {str(e)}"
        
        # Eliminar archivos temporales si existen
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)
        
        # Devolver un mensaje de error como respuesta
        raise HTTPException(
            status_code=500,
            detail=error_message
        )
            
    
    